from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io

RED   = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT  = "Cambria"
SZ    = Pt(11)
HEADER_IMAGE = os.path.join(os.path.dirname(__file__), "header.png")

def _run(para, text, bold=False, color=BLACK, size=None, underline=False):
    run = para.add_run(text)
    run.font.name      = FONT
    run.font.bold      = bold
    run.font.color.rgb = color
    run.font.size      = size or SZ
    run.font.underline = underline
    return run

def _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(14)
    return para

def _red_bottom_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C00000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_right_tab(para, pos_twips):
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def build_ai_notice_docx(ref_no, flat_no, name, issued_date, subject, ai_body_text):
    """
    Wraps AI-generated body text in the society's standard letterhead DOCX format.
    ai_body_text: plain string, paragraphs separated by newlines.
    """
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header image ──────────────────────────────────────────
    if os.path.exists(HEADER_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(HEADER_IMAGE, width=Inches(6.3))

    # ── Red separator ─────────────────────────────────────────
    sep = _para(doc, space_before=2, space_after=6)
    _red_bottom_border(sep)

    # ── Ref No | Date ─────────────────────────────────────────
    p = _para(doc, space_before=6, space_after=6)
    _add_right_tab(p, 9072)
    _run(p, "Ref.No.", bold=True, color=RED, underline=True)
    _run(p, f" {ref_no}", bold=True, color=RED)
    _run(p, "\t", bold=True, color=RED)
    _run(p, f"Date: {issued_date}", bold=True, color=RED)

    _para(doc, space_after=4)

    # ── To section ────────────────────────────────────────────
    p = _para(doc, space_after=2)
    _run(p, "To,")

    if flat_no:
        p = _para(doc, space_after=0)
        _run(p, "Building No & Flat No: ", bold=True)
        _run(p, flat_no, bold=True)

        p = _para(doc, space_after=0)
        _run(p, "Shreeji Iconic CHS Ltd, New Panvel Highway Link Road,", bold=True)

    if name:
        p = _para(doc, space_after=10)
        _run(p, "Member Name: ", bold=True)
        _run(p, name, bold=True)

    _para(doc, space_after=4)

    # ── Subject ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, subject, bold=True, underline=True)

    # ── AI Body paragraphs ────────────────────────────────────
    for line in ai_body_text.strip().split("\n"):
        line = line.strip()
        if not line:
            _para(doc, space_after=4)
            continue
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
        _run(p, line)

    # ── Closing ───────────────────────────────────────────────
    _para(doc, space_after=16)
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Chairman / Secretary", bold=True)
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Shreeji Iconic CHS Ltd.", bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_mom_docx(mom_text, meeting_date, society_name="Shreeji Iconic CHS Ltd."):
    """
    Wraps AI-generated Marathi MOM text into a formatted DOCX.
    """
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ── Header image ──────────────────────────────────────────
    if os.path.exists(HEADER_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(HEADER_IMAGE, width=Inches(6.3))

    sep = _para(doc, space_before=2, space_after=10)
    _red_bottom_border(sep)

    # ── Title ─────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    _run(p, "इतिवृत्त (Minutes of Meeting)", bold=True, size=Pt(14), color=RED)

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _run(p, society_name, bold=True, size=Pt(12))

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    _run(p, f"बैठकीची तारीख: {meeting_date}", bold=True, size=Pt(11))

    sep2 = _para(doc, space_before=2, space_after=12)
    _red_bottom_border(sep2)

    # ── MOM content ───────────────────────────────────────────
    for line in mom_text.strip().split("\n"):
        line = line.strip()
        if not line:
            _para(doc, space_after=4)
            continue
        # Detect section headers (lines ending with : or starting with क्र.)
        is_header = line.endswith(":") or line.startswith("क्र.") or line.startswith("##")
        line = line.lstrip("#").strip()
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT if is_header else WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=6 if is_header else 0, space_after=6)
        run = p.add_run(line)
        run.font.name = "Mangal"   # Devanagari-compatible font
        run.font.bold = is_header
        run.font.color.rgb = RED if is_header else BLACK
        run.font.size = Pt(11)

    # ── Signature section ─────────────────────────────────────
    _para(doc, space_after=20)
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    run = p.add_run("सचिव / अध्यक्ष")
    run.font.name = "Mangal"; run.font.bold = True; run.font.size = Pt(11)

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    run = p.add_run(society_name)
    run.font.name = FONT; run.font.bold = True; run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
