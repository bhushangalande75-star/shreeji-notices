from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io

RED   = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT  = "Cambria"
SZ    = Pt(11)
HEADER_IMAGE = os.path.join(os.path.dirname(__file__), "header.png")

def _run(para, text, bold=False, color=BLACK, size=None, underline=False, superscript=False):
    run = para.add_run(text)
    run.font.name      = FONT
    run.font.bold      = bold
    run.font.color.rgb = color
    run.font.size      = size or SZ
    run.font.underline = underline
    if superscript: run.font.superscript = True
    return run

def _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(14)
    return para

def _red_bottom_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C00000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_right_tab(para, pos_twips):
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)

def generate_notice_3rd(flat_no, ref_no, name, amount,
                        prev_ref_no_1st, prev_ref_no_2nd,
                        issued_date="", due_date="31st March 2026",
                        maintenance_period="March 2026",
                        subject="Sub: 3rd & Final Notice — Legal Action for Recovery of Outstanding Maintenance Dues."):
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header image ──────────────────────────────────────────
    if os.path.exists(HEADER_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(HEADER_IMAGE, width=Inches(6.3))

    # ── Red separator ─────────────────────────────────────────
    sep = _para(doc, space_before=2, space_after=6)
    _red_bottom_border(sep)

    # ── Ref No | Date ─────────────────────────────────────────
    p = _para(doc, space_before=6, space_after=6)
    _add_right_tab(p, 9072)
    _run(p, "Ref.No.", bold=True, color=RED, underline=True)
    _run(p, f" {ref_no}", bold=True, color=RED)
    _run(p, "\t", bold=True, color=RED)
    _run(p, f"Date: {issued_date}", bold=True, color=RED)

    _para(doc, space_after=4)

    # ── To section ────────────────────────────────────────────
    p = _para(doc, space_after=2)
    _run(p, "To,")

    p = _para(doc, space_after=0)
    _run(p, "Building No & Flat No: ", bold=True)
    _run(p, flat_no, bold=True)

    p = _para(doc, space_after=0)
    _run(p, "Shreeji Iconic CHS Ltd, New Panvel Highway Link Road,", bold=True)

    p = _para(doc, space_after=10)
    _run(p, "Member Name: ", bold=True)
    _run(p, name, bold=True)

    _para(doc, space_after=4)

    # ── Subject ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, subject, bold=True, underline=True, color=RED)

    # ── Reference to previous notices ────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "We wish to draw your urgent attention to our previous notices bearing Ref. No. ")
    _run(p, prev_ref_no_1st, bold=True)
    _run(p, " (1st Notice) and Ref. No. ")
    _run(p, prev_ref_no_2nd, bold=True)
    _run(p, " (2nd Notice), which were duly served upon you requesting immediate clearance of your outstanding maintenance dues towards Shreeji Iconic Co-operative Housing Society Ltd. Despite repeated reminders, the same remains unpaid to date.")

    # ── Regret ────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "It is deeply regretted to note that despite the aforesaid notices, you have neither paid the outstanding dues nor have you communicated any valid reason for the default. Your continued non-compliance is not only in violation of the Society's bye-laws but is also causing inconvenience to the proper administration of the Society.")

    # ── Outstanding amount ────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "As per the Society's records, the total outstanding maintenance dues payable by you are as follows:")

    formatted_amount = f"{amount:,}"
    p = _para(doc, space_before=4, space_after=8)
    _run(p, f"\u27A4  Total Outstanding Maintenance Charges till {maintenance_period}: - Rs. ", bold=True, color=RED)
    _run(p, formatted_amount, bold=True, color=RED)

    # ── Final demand ──────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "This is therefore the ")
    _run(p, "3rd and FINAL NOTICE", bold=True, underline=True)
    _run(p, f" served upon you demanding that you forthwith pay the total outstanding amount of Rs. {formatted_amount}/- on or before ")
    _run(p, due_date, bold=True)
    _run(p, " failing which the Managing Committee shall be constrained to initiate the following legal proceedings without any further notice:")

    # ── Legal consequences bullet ─────────────────────────────
    consequences = [
        "Filing of Recovery Application u/s 154(B)(29) of the Maharashtra Co-operative Societies Act, 1960 before the Assistant / Deputy Registrar, Co-operative Societies, Ambernath for recovery of dues as arrears of Land Revenue.",
        "Recovery of all legal expenses, costs, and charges incurred by the Society in connection with the above proceedings shall be borne exclusively by the defaulting member.",
        "Restriction of access to Society amenities and common facilities as per the Society's bye-laws until full clearance of dues.",
        "Any further action as deemed fit and proper under applicable laws and bye-laws of the Society.",
    ]
    for c in consequences:
        p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=2, space_after=6)
        _run(p, "\u27A4  ", bold=True, color=RED)
        _run(p, c, bold=False)

    # ── Warning ───────────────────────────────────────────────
    _para(doc, space_after=4)
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "Please be advised that the legal proceedings once initiated shall not be withdrawn and all costs including advocate fees, court fees, and incidental expenses shall be recovered from you in addition to the principal dues. You are therefore earnestly advised to avoid such an unpleasant situation by making the payment immediately.")

    # ── Closing ───────────────────────────────────────────────
    p = _para(doc, space_after=24)
    _run(p, "No further communication shall be entertained in this regard.")

    # ── Signature ────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Chairman / Secretary", bold=True)

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Shreeji Iconic CHS Ltd.", bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
