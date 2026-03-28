from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io

RED   = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT  = "Cambria"
SZ    = Pt(11)

HEADER_IMAGE = os.path.join(os.path.dirname(__file__), "header.png")

def _run(para, text, bold=False, color=BLACK, size=None, underline=False, superscript=False):
    run = para.add_run(text)
    run.font.name      = FONT
    run.font.bold      = bold
    run.font.color.rgb = color
    run.font.size      = size or SZ
    run.font.underline = underline
    if superscript:
        run.font.superscript = True
    return run

def _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(14)
    return para

def _red_bottom_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C00000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_right_tab(para, pos_twips):
    """Add a right-aligned tab stop to a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)

def generate_notice(flat_no, ref_no, name, amount):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header image ──────────────────────────────────────────
    if os.path.exists(HEADER_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(HEADER_IMAGE, width=Inches(6.3))

    # ── Red separator line ────────────────────────────────────
    sep = _para(doc, space_before=2, space_after=6)
    _red_bottom_border(sep)

    # ── Ref No | Date — single paragraph, tab to right ───────
    p = _para(doc, space_before=6, space_after=6)
    # Right tab at ~6.3 inches (content width) = 9072 twips
    _add_right_tab(p, 9072)
    _run(p, "Ref.No.", bold=True, color=RED, underline=True)
    _run(p, f"{ref_no}", bold=True, color=RED)
    _run(p, "\t", bold=True, color=RED)                                          # jump to right edge
    _run(p, "Date: 23/03/2026", bold=True, color=RED)

    # ── Spacer ────────────────────────────────────────────────
    _para(doc, space_after=4)

    # ── To, ──────────────────────────────────────────────────
    p = _para(doc, space_after=2)
    _run(p, "To,")

    # Building No & Flat No
    p = _para(doc, space_after=0)
    _run(p, "Building No & Flat No: ", bold=True)
    _run(p, flat_no, bold=True)

    # Address
    p = _para(doc, space_after=0)
    _run(p, "Shreeji Iconic CHS Ltd, New Panvel Highway Link Road,", bold=True)

    # Member Name
    p = _para(doc, space_after=10)
    _run(p, "Member Name: ", bold=True)
    _run(p, name, bold=True)

    # ── Spacer ────────────────────────────────────────────────
    _para(doc, space_after=4)

    # ── Subject ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, "Sub: Notice for Recovery of Due Maintenance.", bold=True, underline=True)

    # ── Body 1 ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "You are well aware that as per the provisions of our Societies\u2019 bye laws you need to pay your outstanding dues regularly within prescribed period. This notice is being served to you as per resolution passed in SGM held on 15")
    _run(p, "th", superscript=True)
    _run(p, " October 2023 to send notices to defaulters who has not paid society maintenance for a minimum of 3 months as per provisions of law on date 17")
    _run(p, "th", superscript=True)
    _run(p, " June 2024.")

    # ── Body 2 ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "It has been observed from the society record that the total outstanding towards the society\u2019s contribution is receivable by the society from you which is as under:")

    # ── Bullet ───────────────────────────────────────────────
    formatted_amount = f"{amount:,}"
    p = _para(doc, space_before=4, space_after=8)
    _run(p, "\u27A4  Maintenance Charges: - Rs. ", bold=True)
    _run(p, formatted_amount, bold=True)

    # ── Body 3 ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "Therefore, you are requested to make an immediate payment of above amount by dated: - 31")
    _run(p, "st", superscript=True)
    _run(p, " March 2026 to avoid unpleasant situation for the committee in adherence of legal provision of law and filing recovery application u/sec 154 (B) 29 of MCS Act, 1960 to read with several provisions of laws and bye-laws.")

    # ── Body 4 ───────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "Please note that if you force the society to initiate legal action as above or others, the application for recovery of dues will be filed against you at Asst. Registrar / Deputy Registrar of Co-op. Societies, Ambernath, for issue of recovery certificated u/s 154B (29) of the Maharashtra Co-operative Societies Act.1960 for recovery of dues as arrears of LAND REVENUE and under the said resolution we have been authorized to RECOVER LEGAL EXPENCES from you to initiate further legal action.")

    # ── Appreciation ─────────────────────────────────────────
    p = _para(doc, space_after=24)
    _run(p, "Your priority to this will be appreciated.")

    # ── Signature ────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Chairman / Secretary")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
