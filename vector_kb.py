"""
vector_kb.py — Document processing pipeline for SocietyNotice Pro Vector KB.

Responsibilities:
  1. Extract text from PDF / DOCX / XLSX / TXT
  2. Chunk text into overlapping segments
  3. Embed chunks using fastembed (ONNX-based, NO torch required)
  4. Store in Neon pgvector

Embedding model: sentence-transformers/all-MiniLM-L6-v2 via fastembed
  - Dimension: 384  (same as before — no DB migration needed)
  - Uses ONNX runtime instead of PyTorch → ~80MB RAM vs ~600MB
  - Zero API key, runs fully local
"""

import io, re

EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
EMBED_DIM   = 384

# Lazy-load the fastembed model
_embedder = None

def _get_embedder():
    global _embedder
    if _embedder is None:
        from fastembed import TextEmbedding
        print(f"[KB] Loading fastembed model {EMBED_MODEL}...")
        _embedder = TextEmbedding(model_name=EMBED_MODEL)
        print("[KB] Model loaded ✅")
    return _embedder


CHUNK_SIZE    = 500  # characters
CHUNK_OVERLAP = 80


# ── Text Extraction ────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t.strip())
        return "\n\n".join(pages)
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(file_bytes))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paras.append(row_text)
        return "\n\n".join(paras)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")


def extract_text_from_excel(file_bytes: bytes) -> str:
    try:
        import pandas as pd
        dfs      = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        sections = []
        for sheet_name, df in dfs.items():
            df   = df.fillna("").astype(str)
            rows = [" | ".join(str(c) for c in df.columns), "-" * 60]
            for _, row in df.iterrows():
                line = " | ".join(str(v).strip() for v in row.values)
                if any(v.strip() for v in row.values if v.strip() not in ("nan", "")):
                    rows.append(line)
            sections.append(f"Sheet: {sheet_name}\n" + "\n".join(rows))
        return "\n\n".join(sections)
    except Exception as e:
        raise ValueError(f"Excel extraction failed: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return extract_text_from_excel(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


# ── Chunking ───────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE,
               overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    chunks = []
    start  = 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        para_break = text.rfind("\n\n", start, end)
        if para_break > start + overlap:
            end = para_break
        else:
            sent_break = max(
                text.rfind(". ",  start + overlap, end),
                text.rfind("। ", start + overlap, end),
                text.rfind("\n",  start + overlap, end),
            )
            if sent_break > start + overlap:
                end = sent_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap

    return chunks


# ── Embedding ──────────────────────────────────────────────────

def embed_texts(texts: list[str]) -> list[list[float]]:
    """
    Embed a list of texts using fastembed (ONNX, no torch).
    Returns list of 384-dim float vectors.
    """
    if not texts:
        return []
    model = _get_embedder()
    # fastembed.embed() returns a generator of numpy arrays
    return [vec.tolist() for vec in model.embed(texts)]


def embed_query(query: str) -> list[float]:
    """Embed a single query string."""
    model = _get_embedder()
    return next(model.embed([query])).tolist()


# ── Full pipeline ──────────────────────────────────────────────

def process_document(file_bytes: bytes,
                     filename: str) -> tuple[list[str], list[list[float]]]:
    """extract → chunk → embed. Returns (chunks, embeddings)."""
    text   = extract_text(file_bytes, filename)
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("No text could be extracted from this document.")
    embeddings = embed_texts(chunks)
    return chunks, embeddings
