from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io

RED   = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT  = "Cambria"
SZ    = Pt(11)
HEADER_IMAGE = os.path.join(os.path.dirname(__file__), "header.png")

def _run(para, text, bold=False, color=BLACK, size=None, underline=False, superscript=False):
    run = para.add_run(text)
    run.font.name      = FONT
    run.font.bold      = bold
    run.font.color.rgb = color
    run.font.size      = size or SZ
    run.font.underline = underline
    if superscript: run.font.superscript = True
    return run

def _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=8):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(14)
    return para

def _red_bottom_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C00000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_right_tab(para, pos_twips):
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)

def generate_notice_2nd(flat_no, ref_no, name, amount, prev_ref_no, issued_date="", due_date="31st March 2026"):
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Header image
    if os.path.exists(HEADER_IMAGE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run().add_picture(HEADER_IMAGE, width=Inches(6.3))

    # Red separator
    sep = _para(doc, space_before=2, space_after=6)
    _red_bottom_border(sep)

    # Ref No | Date
    p = _para(doc, space_before=6, space_after=6)
    _add_right_tab(p, 9072)
    _run(p, "Ref.No.", bold=True, color=RED, underline=True)
    _run(p, f" {ref_no}", bold=True, color=RED)
    _run(p, "\t", bold=True, color=RED)
    _run(p, f"Date: {issued_date}", bold=True, color=RED)

    _para(doc, space_after=4)

    # To section
    p = _para(doc, space_after=2)
    _run(p, "To,")

    p = _para(doc, space_after=0)
    _run(p, "Building No & Flat No: ", bold=True)
    _run(p, flat_no, bold=True)

    p = _para(doc, space_after=0)
    _run(p, "Shreeji Iconic CHS Ltd, New Panvel Highway Link Road,", bold=True)

    p = _para(doc, space_after=10)
    _run(p, "Member Name: ", bold=True)
    _run(p, name, bold=True)

    _para(doc, space_after=4)

    # Subject
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, "Sub: 2nd Notice for Recovery of Outstanding Maintenance Dues.", bold=True, underline=True)

    # Dear Member
    p = _para(doc, space_after=8)
    _run(p, "Dear Member,")

    # Body 1 - Reference to 1st notice
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "We wish to draw your kind attention to our earlier notice bearing ")
    _run(p, f"Ref. No. {prev_ref_no}", bold=True)
    _run(p, ", wherein you were requested to clear your outstanding maintenance dues towards Shreeji Iconic Co-operative Housing Society Ltd.")

    # Body 2 - Outstanding dues
    formatted_amount = f"{amount:,}"
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "It is observed from the society\u2019s records that despite the said notice, the outstanding amount of ")
    _run(p, f"Rs. {formatted_amount}/-", bold=True)
    _run(p, f" towards maintenance charges till {due_date} remains unpaid as on date.")

    # Bullet
    p = _para(doc, space_before=4, space_after=8)
    _run(p, "\u27A4  Maintenance Charges till March 2026: - Rs. ", bold=True)
    _run(p, formatted_amount, bold=True)

    # Body 3 - Request
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "We therefore once again ")
    _run(p, "earnestly request", bold=True)
    _run(p, " you to make the payment of the outstanding amount at the earliest. As a law-abiding society, we are bound by the provisions of the ")
    _run(p, "Maharashtra Co-operative Societies Act, 1960", bold=True)
    _run(p, " and the bye-laws of our society to take necessary steps for recovery of dues.")

    # Body 4 - Warning
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _run(p, "Please note that failure to make the payment will constrain the Managing Committee to initiate recovery proceedings against you under ")
    _run(p, "Section 154(B)(29) of the MCS Act, 1960", bold=True)
    _run(p, ", including filing of a recovery application before the ")
    _run(p, "Asst. Registrar / Deputy Registrar of Co-operative Societies, Ambernath", bold=True)
    _run(p, ", for recovery of dues as arrears of ")
    _run(p, "Land Revenue", bold=True)
    _run(p, ". The legal expenses incurred thereof shall also be recovered from you.")

    # Closing
    p = _para(doc, space_after=8)
    _run(p, "We sincerely hope that you will give this matter your ")
    _run(p, "immediate attention", bold=True)
    _run(p, " and cooperate with the society to avoid any legal action.")

    p = _para(doc, space_after=24)
    _run(p, "Your priority to this will be greatly appreciated.")

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _run(p, "Chairman / Secretary")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
